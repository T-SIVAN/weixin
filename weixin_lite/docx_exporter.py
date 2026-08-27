"""Native, editable Word export for a generated article.

The Markdown remains the source of truth, but it is mapped to Word paragraphs,
headings, lists and inline images instead of being wrapped in an HTML document.
"""
from __future__ import annotations

import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import PurePosixPath
from urllib.parse import unquote

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from PIL import Image, UnidentifiedImageError

from .models import EditableTable, QuickReadArticle


_BOLD_RE = re.compile(r"(\*\*.+?\*\*)")
_HEADING_RE = re.compile(r"^(#{1,6})\s+(.+?)\s*#*\s*$")
_OPTIONAL_IMAGE_TITLE_RE = re.compile(r"(?:\"(?:\\.|[^\"\\])*\"|'(?:\\.|[^'\\])*'|\((?:\\.|[^)\\])*\))$")


@dataclass(frozen=True)
class _MarkdownImageReference:
    alt: str
    name: str | None
    standalone: bool
    line_number: int


class DocxExportError(ValueError):
    """Raised when an editable Word document would be missing a referenced image."""


def _local_image_name(destination: str) -> str | None:
    value = unquote(destination.strip().strip("<>").replace("\\", "/"))
    value = value.split("?", 1)[0].split("#", 1)[0]
    for prefix in ("images/", "./images/", "/images/"):
        if value.startswith(prefix):
            name = value[len(prefix) :]
            path = PurePosixPath(name)
            if name and not path.is_absolute() and ".." not in path.parts:
                return path.as_posix()
    return None


def _find_image_closing_parenthesis(line: str, content_start: int) -> int | None:
    depth = 1
    escaped = False
    in_angle_destination = False
    for index in range(content_start, len(line)):
        char = line[index]
        if escaped:
            escaped = False
            continue
        if char == "\\":
            escaped = True
            continue
        if char == "<" and not in_angle_destination:
            in_angle_destination = True
            continue
        if char == ">" and in_angle_destination:
            in_angle_destination = False
            continue
        if in_angle_destination:
            continue
        if char == "(":
            depth += 1
        elif char == ")":
            depth -= 1
            if depth == 0:
                return index
    return None


def _image_destination(content: str) -> str | None:
    """Read a CommonMark image destination, allowing an optional title."""
    value = content.strip()
    if not value:
        return None
    if value.startswith("<"):
        closing = value.find(">")
        if closing < 0:
            return None
        destination = value[1:closing].strip()
        remainder = value[closing + 1 :].strip()
        return destination if not remainder or _OPTIONAL_IMAGE_TITLE_RE.fullmatch(remainder) else None

    title_match = re.search(r"\s+(?:\"(?:\\.|[^\"\\])*\"|'(?:\\.|[^'\\])*'|\((?:\\.|[^)\\])*\))$", value)
    destination = value[: title_match.start()].rstrip() if title_match else value
    return destination or None


def _markdown_image_references(markdown: str) -> list[_MarkdownImageReference]:
    references: list[_MarkdownImageReference] = []
    for line_number, raw_line in enumerate((markdown or "").splitlines(), start=1):
        cursor = 0
        while True:
            image_start = raw_line.find("![", cursor)
            if image_start < 0:
                break
            alt_end = raw_line.find("](", image_start + 2)
            if alt_end < 0:
                break
            image_end = _find_image_closing_parenthesis(raw_line, alt_end + 2)
            if image_end is None:
                cursor = alt_end + 2
                continue
            destination = _image_destination(raw_line[alt_end + 2 : image_end])
            if destination is not None:
                references.append(
                    _MarkdownImageReference(
                        alt=raw_line[image_start + 2 : alt_end],
                        name=_local_image_name(destination),
                        standalone=not raw_line[:image_start].strip() and not raw_line[image_end + 1 :].strip(),
                        line_number=line_number,
                    )
                )
            cursor = image_end + 1
    return references


def markdown_image_names(markdown: str) -> list[str]:
    return [reference.name for reference in _markdown_image_references(markdown) if reference.name]


def _validate_standalone_image_references(markdown: str) -> None:
    inline_names = [reference.name for reference in _markdown_image_references(markdown) if reference.name and not reference.standalone]
    if inline_names:
        names = list(dict.fromkeys(inline_names))
        raise DocxExportError("无法导出 Word：Markdown 图片必须独占一行，内联图片：" + "、".join(names))


def _standalone_image_reference(line: str) -> _MarkdownImageReference | None:
    references = _markdown_image_references(line)
    return references[0] if len(references) == 1 and references[0].standalone else None


def expected_article_image_names(article: QuickReadArticle) -> list[str]:
    """Return the ordered local assets required to render the article faithfully."""
    names = markdown_image_names(article.body_markdown)
    lead = str(article.lead_image_name or "").strip()
    if lead and lead not in names:
        # The HTML renderer injects a missing lead image before the body; Word
        # mirrors that behavior so the two final outputs retain one lead image.
        names.insert(0, lead)
    return list(dict.fromkeys(names))


def validate_article_images(article: QuickReadArticle, image_assets: dict[str, bytes] | None) -> list[str]:
    expected = expected_article_image_names(article)
    assets = image_assets or {}
    missing = [name for name in expected if not assets.get(name)]
    if missing:
        raise DocxExportError("无法导出 Word：正文引用的图片缺失：" + "、".join(missing))
    return expected


def _set_font(run, *, size: float, bold: bool = False, color: str = "000000") -> None:
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run._element.rPr.rFonts.set(qn("w:ascii"), "Aptos")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Aptos")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def _configure_document(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.6)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.65)
    section.right_margin = Cm(1.65)

    normal = document.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(13.5)
    normal.paragraph_format.space_after = Pt(10)
    normal.paragraph_format.line_spacing = 1.8

    for style_name, size, before, after in (("Heading 1", 21, 26, 12), ("Heading 2", 16, 20, 8), ("Heading 3", 14, 14, 6)):
        style = document.styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.45


def _add_rich_text(paragraph, text: str, *, size: float = 13.5) -> None:
    for part in _BOLD_RE.split(text or ""):
        if not part:
            continue
        strong = part.startswith("**") and part.endswith("**")
        run = paragraph.add_run(part[2:-2] if strong else part)
        _set_font(run, size=size, bold=strong)


def _figure_metadata(article: QuickReadArticle, image_name: str, alt: str) -> str:
    if image_name == article.lead_image_name:
        return "论文首页题录截图"
    for figure in article.figures:
        if figure.image_name == image_name:
            parts = [figure.figure_id, f"p.{figure.page}" if figure.page else "", figure.caption]
            return " | ".join(part for part in parts if part)[:512]
    return alt or image_name


def _add_image(document: Document, article: QuickReadArticle, image_name: str, alt: str, image_assets: dict[str, bytes]) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(14)
    paragraph.paragraph_format.space_after = Pt(12)
    run = paragraph.add_run()
    inline_shape = run.add_picture(BytesIO(_word_image_bytes(image_assets[image_name])), width=Cm(14.66))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", _figure_metadata(article, image_name, alt))
    doc_pr.set("title", alt or image_name)


def _add_editable_table(document: Document, table: EditableTable | None) -> None:
    if not table or not table.headers or not table.rows or table.confidence < 0.7:
        return
    native = document.add_table(rows=1, cols=len(table.headers))
    native.style = "Table Grid"
    for index, value in enumerate(table.headers):
        native.rows[0].cells[index].text = value
        for run in native.rows[0].cells[index].paragraphs[0].runs:
            _set_font(run, size=10.5, bold=True)
    for values in table.rows:
        row = native.add_row()
        for index, value in enumerate(values[: len(table.headers)]):
            row.cells[index].text = value
            for run in row.cells[index].paragraphs[0].runs:
                _set_font(run, size=10.5)
    document.add_paragraph()


def _editable_table_for_image(article: QuickReadArticle, image_name: str) -> EditableTable | None:
    for figure in article.figures:
        if figure.image_name == image_name:
            return figure.editable_table
    return None


def _word_image_bytes(data: bytes) -> bytes:
    """Return bytes Word can embed, converting only unsupported image formats."""
    try:
        with Image.open(BytesIO(data)) as source:
            image_format = (source.format or "").upper()
            if image_format in {"PNG", "JPEG", "GIF", "BMP"}:
                return data
            image = source.copy()
    except (UnidentifiedImageError, OSError) as exc:
        raise DocxExportError("无法读取正文图片，无法嵌入 Word。") from exc

    output = BytesIO()
    # PNG preserves alpha for screenshots and is widely supported by Word/WPS.
    if image.mode not in {"RGB", "RGBA"}:
        image = image.convert("RGBA" if "transparency" in image.info else "RGB")
    image.save(output, format="PNG")
    return output.getvalue()


def export_article_docx(article: QuickReadArticle, image_assets: dict[str, bytes] | None) -> bytes:
    """Build an editable DOCX with all article images embedded in the package."""
    expected = validate_article_images(article, image_assets)
    _validate_standalone_image_references(article.body_markdown)
    assets = image_assets or {}
    document = Document()
    _configure_document(document)

    title = document.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    title.paragraph_format.space_after = Pt(14)
    _add_rich_text(title, article.title, size=22)
    for run in title.runs:
        run.font.bold = True

    lines = list((article.body_markdown or "").splitlines())
    lead = str(article.lead_image_name or "").strip()
    markdown_names = markdown_image_names(article.body_markdown)
    if lead and lead in expected and lead not in markdown_names:
        _add_image(document, article, lead, "论文首页", assets)

    inserted_images: set[str] = set()
    if lead and lead in expected and lead not in markdown_names:
        inserted_images.add(lead)

    for raw_line in lines:
        line = raw_line.strip()
        if not line:
            continue
        image_reference = _standalone_image_reference(raw_line)
        if image_reference and image_reference.name:
            name = image_reference.name
            # The article title image is allowed only once even if Markdown has
            # an accidental duplicate reference.
            if name and not (name == lead and name in inserted_images):
                _add_image(document, article, name, image_reference.alt, assets)
                _add_editable_table(document, _editable_table_for_image(article, name))
                inserted_images.add(name)
            continue
        heading_match = _HEADING_RE.match(line)
        if heading_match:
            level = len(heading_match.group(1))
            heading = heading_match.group(2)
            # Generated Markdown usually starts with the same title that has
            # already been rendered as the document title.
            if level == 1 and heading == article.title:
                continue
            style_name, size = (
                ("Heading 1", 21) if level <= 2 else ("Heading 2", 16) if level == 3 else ("Heading 3", 14)
            )
            paragraph = document.add_paragraph(style=style_name)
            _add_rich_text(paragraph, heading, size=size)
            continue
        if line.startswith("> "):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.left_indent = Cm(0.5)
            paragraph.paragraph_format.space_after = Pt(8)
            _add_rich_text(paragraph, line[2:], size=11.5)
            continue
        number_match = re.match(r"^(\d+)\.\s+(.+)$", line)
        if number_match:
            paragraph = document.add_paragraph(style="List Number")
            _add_rich_text(paragraph, number_match.group(2))
            continue
        if line.startswith(("- ", "* ")):
            paragraph = document.add_paragraph(style="List Bullet")
            _add_rich_text(paragraph, line[2:])
            continue
        if line.startswith("**") and line.endswith("**"):
            paragraph = document.add_paragraph()
            paragraph.paragraph_format.space_before = Pt(12)
            paragraph.paragraph_format.space_after = Pt(7)
            _add_rich_text(paragraph, line, size=15)
            continue
        paragraph = document.add_paragraph()
        _add_rich_text(paragraph, line)

    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()
