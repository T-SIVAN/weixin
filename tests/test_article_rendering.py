from __future__ import annotations

import io
import zipfile

import pytest
from PIL import Image

from weixin_lite.docx_exporter import DocxExportError
from weixin_lite.exporter import article_document_html, article_html, export_article_docx_bytes, export_article_html, project_zip, wechat_content_html
from weixin_lite.models import BatchProject, FigureAnalysis, PaperInput, QuickReadArticle
from weixin_lite.wechat_publish import (
    WechatDraftConfig,
    WechatPublishError,
    build_draft_payload,
    publish_draft,
    replace_content_image_sources,
)


def make_article(body_html: str = "<p style=\"margin:20px 0;font-size:18px;line-height:2.05;\">正文</p>") -> QuickReadArticle:
    return QuickReadArticle(
        paper=PaperInput(title_zh="论文题名", url="https://doi.org/10.1000/example"),
        title="平台文章标题",
        digest="平台摘要",
        body_markdown="正文",
        body_html=body_html,
        cover_image_name="platform-cover.png",
    )


def png_bytes(color: str = "white") -> bytes:
    buffer = io.BytesIO()
    Image.new("RGB", (24, 16), color=color).save(buffer, format="PNG")
    return buffer.getvalue()


def test_preview_export_and_draft_share_one_inline_body_without_platform_header():
    article = make_article(
        '<p style="margin:44px 0 22px;font-size:22px;line-height:1.45;font-weight:800;">章节标题</p>'
        '<p style="margin:30px 0 18px;font-size:20px;line-height:1.55;font-weight:800;">图文小标题</p>'
        '<p style="margin:20px 0;font-size:18px;line-height:2.05;">正文</p>'
    )
    config = WechatDraftConfig(author="公众号作者")

    body = article_document_html(article, author="不应进入正文", account_name="不应进入正文")
    exported = article_html(article)
    payload = build_draft_payload(article, config)

    assert body in exported
    assert payload["articles"][0]["content"] == body
    assert payload["articles"][0]["title"] == article.title
    assert payload["articles"][0]["digest"] == article.digest
    assert payload["articles"][0]["author"] == "公众号作者"
    assert "原创" not in body
    assert "公众号作者" not in body
    assert "不应进入正文" not in body
    assert article.title not in body
    assert "platform-cover.png" not in body
    assert "max-width:760px" in body
    assert "font-size:18px;line-height:2.05" in body
    assert "font-size:22px" in body
    assert "font-size:20px" in body


def test_lead_image_is_inserted_once_and_cover_is_never_inserted():
    article = make_article('<p>导语</p><img src="images/paper-lead.png" alt="论文首页">')
    article.lead_image_name = "paper-lead.png"

    body = article_document_html(article)

    assert body.count('src="images/paper-lead.png"') == 1
    assert "platform-cover.png" not in body

    article.body_html = "<p>导语</p>"
    body_without_existing_lead = wechat_content_html(article, include_cover=True)
    assert body_without_existing_lead.count('src="images/paper-lead.png"') == 1
    assert "platform-cover.png" not in body_without_existing_lead


def test_dry_run_reports_missing_body_images_and_real_publish_fails_before_api(monkeypatch):
    article = make_article('<p>图解</p><img src="images/missing-figure.png" alt="Fig. 1">')
    config = WechatDraftConfig(cover_image_name="platform-cover.png")
    assets = {"platform-cover.png": b"cover"}

    dry_run = publish_draft(article, config, image_assets=assets, dry_run=True)

    assert dry_run["referenced_content_images"] == ["missing-figure.png"]
    assert any("missing-figure.png" in item for item in dry_run["diagnostics"])

    monkeypatch.setattr(
        "weixin_lite.wechat_publish.get_access_token",
        lambda *_args: pytest.fail("缺图时不应请求微信 access token"),
    )
    with pytest.raises(WechatPublishError, match="missing-figure.png"):
        publish_draft(article, config, image_assets=assets, dry_run=False)


def test_content_upload_only_uses_referenced_assets_and_deduplicates(monkeypatch):
    uploaded: list[str] = []

    def fake_upload(_token: str, image_name: str, _data: bytes) -> str:
        uploaded.append(image_name)
        return f"https://mmbiz.example/{image_name}"

    monkeypatch.setattr("weixin_lite.wechat_publish.upload_content_image", fake_upload)
    content = (
        '<img src="images/used.png"><p>正文</p>'
        "<img src='images/used.png'><img src=\"https://example.com/external.png\">"
    )
    assets = {"used.png": b"used", "unused.png": b"unused"}

    replaced = replace_content_image_sources(content, "token", assets)

    assert uploaded == ["used.png"]
    assert replaced.count("https://mmbiz.example/used.png") == 2
    assert "unused.png" not in replaced


def test_project_zip_preserves_unicode_spaces_and_rewrites_colliding_asset_names():
    article = make_article(
        '<img src="images/组甲/结果 图.png"><img src="images/组乙/结果 图.png">'
    )
    article.body_markdown = "![甲](images/组甲/结果 图.png)\n![乙](<images/组乙/结果 图.png>)"
    project = BatchProject(topic="测试项目", articles=[article])

    archive = project_zip(
        project,
        {"组甲/结果 图.png": png_bytes("red"), "组乙/结果 图.png": png_bytes("blue"), "../危险.png": png_bytes("green")},
    )

    with zipfile.ZipFile(io.BytesIO(archive)) as zf:
        names = zf.namelist()
        html_name = next(name for name in names if name.endswith(".html"))
        markdown_name = next(name for name in names if name.endswith(".md"))
        rendered_html = zf.read(html_name).decode("utf-8")
        markdown = zf.read(markdown_name).decode("utf-8")

        assert "images/结果 图.png" in names
        assert "images/结果 图-2.png" in names
        assert "images/危险.png" in names
        assert not any(".." in name for name in names)
        assert 'src="images/结果 图.png"' in rendered_html
        assert 'src="images/结果 图-2.png"' in rendered_html
        assert "images/结果 图.png" in markdown
        assert "images/结果 图-2.png" in markdown
        assert "组甲/" not in rendered_html + markdown
        assert "组乙/" not in rendered_html + markdown


def test_editable_docx_embeds_lead_and_confirmed_figure_once():
    article = make_article()
    article.title = "可编辑论文解读"
    article.lead_image_name = "lead.png"
    article.body_markdown = """# 可编辑论文解读

![论文首页](images/lead.png)

导语中的**关键术语**可编辑。

## 关键图证据解读

![Fig. 2](images/fig2.png)

**Fig. 2：关键结果图**

图展示什么：验证结果。
"""
    article.figures = [FigureAnalysis("Fig. 2", "结果图", page="4", image_name="fig2.png", selected=True)]
    data = export_article_docx_bytes(article, {"lead.png": png_bytes(), "fig2.png": png_bytes("blue")})

    from docx import Document

    document = Document(io.BytesIO(data))
    assert "可编辑论文解读" in "\n".join(item.text for item in document.paragraphs)
    assert "关键术语" in "\n".join(item.text for item in document.paragraphs)
    with zipfile.ZipFile(io.BytesIO(data)) as archive:
        images = [name for name in archive.namelist() if name.startswith("word/media/")]
    assert len(images) == 2


def test_docx_export_rejects_missing_referenced_image_and_portable_html_embeds_bytes():
    article = make_article('<img src="images/lead.png" alt="论文首页">')
    article.body_markdown = "# 标题\n\n![论文首页](images/lead.png)"
    article.lead_image_name = "lead.png"

    with pytest.raises(DocxExportError, match="lead.png"):
        export_article_docx_bytes(article, {})

    portable = export_article_html(article, {"lead.png": png_bytes()}).decode("utf-8")
    assert "data:image/png;base64," in portable


def test_src_parsing_normalizes_local_forms_preserves_alt_and_ignores_data_src(monkeypatch):
    uploaded: list[str] = []

    def fake_upload(_token: str, image_name: str, _data: bytes) -> str:
        uploaded.append(image_name)
        return f"https://mmbiz.example/{image_name.replace(' ', '%20')}"

    monkeypatch.setattr("weixin_lite.wechat_publish.upload_content_image", fake_upload)
    content = (
        '<img alt="images/图 一.png" data-src="images/lazy.png" src="./images/图 一.png">'
        "<img src='/images/图二.png'><img data-src='images/ignored.png'>"
        '<img src="https://example.com/external.png">'
    )

    replaced = replace_content_image_sources(content, "token", {"图 一.png": b"one", "图二.png": b"two"})

    assert uploaded == ["图 一.png", "图二.png"]
    assert 'alt="images/图 一.png"' in replaced
    assert 'data-src="images/lazy.png"' in replaced
    assert "data-src='images/ignored.png'" in replaced
    assert 'src="https://mmbiz.example/图%20一.png"' in replaced
    assert "src='https://mmbiz.example/图二.png'" in replaced
    assert 'src="https://example.com/external.png"' in replaced


@pytest.mark.parametrize("bad_src", [r"C:\\temp\\figure.png", "file:///tmp/figure.png", "/tmp/figure.png", "images/../secret.png"])
def test_invalid_local_image_sources_are_rejected(bad_src):
    article = make_article(f'<img src="{bad_src}">')
    with pytest.raises(WechatPublishError):
        publish_draft(article, WechatDraftConfig(), dry_run=True)


def test_dry_run_reports_external_images():
    article = make_article('<img src="https://example.com/external.png">')
    result = publish_draft(article, WechatDraftConfig(), dry_run=True)

    assert result["external_images"] == ["https://example.com/external.png"]
    assert any("外链图片" in item for item in result["diagnostics"])


def test_real_publish_uploads_only_referenced_images_and_builds_final_payload(monkeypatch):
    article = make_article(
        '<img alt="images/used.png" src="images/used.png">'
        '<img src="./images/second image.png"><img src="https://example.com/keep.png">'
    )
    config = WechatDraftConfig(app_id="app", app_secret="secret", cover_image_name="platform-cover.png")
    calls: dict[str, object] = {"content": []}

    monkeypatch.setattr("weixin_lite.wechat_publish.get_access_token", lambda *_args: "token")
    monkeypatch.setattr("weixin_lite.wechat_publish.upload_cover_material", lambda *_args: "cover-media-id")

    def fake_content_upload(_token: str, name: str, _data: bytes) -> str:
        calls["content"].append(name)
        return "https://mmbiz.example/" + urllib_quote(name)

    def fake_create(_token: str, payload: dict):
        calls["payload"] = payload
        return {"media_id": "draft-id"}

    monkeypatch.setattr("weixin_lite.wechat_publish.upload_content_image", fake_content_upload)
    monkeypatch.setattr("weixin_lite.wechat_publish.create_draft", fake_create)

    result = publish_draft(
        article,
        config,
        image_assets={
            "platform-cover.png": b"cover",
            "used.png": b"used",
            "second image.png": b"second",
            "unused.png": b"unused",
        },
        dry_run=False,
    )

    payload = result["payload"]
    content = payload["articles"][0]["content"]
    assert calls["content"] == ["used.png", "second image.png"]
    assert payload is calls["payload"]
    assert payload["articles"][0]["thumb_media_id"] == "cover-media-id"
    assert 'alt="images/used.png"' in content
    assert 'src="https://mmbiz.example/used.png"' in content
    assert 'src="https://mmbiz.example/second%20image.png"' in content
    assert 'src="https://example.com/keep.png"' in content
    assert "unused.png" not in content
    assert result["result"] == {"media_id": "draft-id"}


def urllib_quote(value: str) -> str:
    from urllib.parse import quote

    return quote(value)
